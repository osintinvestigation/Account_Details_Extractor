"""
GOOGLE ZIP CONSOLIDATED INVESTIGATION REPORT GENERATOR
======================================================

Input:
    Google export ZIP supplied by the user.

Output:
    1. Consolidated_Google_Investigation_Report.xlsx
    2. Consolidated_Google_Investigation_Report.docx
    3. IP_Details_Formatted.csv
    4. 2405_2409_IP_Details.docx
    5. Remaining_IP_Details.docx
    6. 2405_2409_IP_Details.txt
    7. Remaining_IP_Details.txt

Important IP formatting rule:
    - 2405: / 2409: IPv6:
          YYYYMMDD HHMMSS
    - 2401: IPv6 (Airtel):
          DD/MMM/YYYY HH:MM:SS
    - All other IP addresses:
          DD-MMM-YYYY HH:MM:SS

Duplicate IP addresses:
    - Removed from the formatted IP Details output.
    - First occurrence is retained.

The script automatically handles nested ZIP files. If the Google export has no explicit F DATE/F TIME/T DATE/T TIME fields, the Google Subscriber IP Activity timestamp is used as FROM and TO for the separate IP Word files.
"""

import os
import sys
import shutil
import zipfile
from datetime import datetime, timezone, timedelta

import pandas as pd
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


# ============================================================
# USER INPUT
# ============================================================

def get_input_paths():
    print("=" * 75)
    print(" GOOGLE CONSOLIDATED INVESTIGATION REPORT GENERATOR")
    print("=" * 75)

    while True:
        zip_path = input(
            "\nEnter FULL path of Google ZIP file:\n> "
        ).strip().strip('"')

        if (
            os.path.isfile(zip_path)
            and zipfile.is_zipfile(zip_path)
        ):
            break

        print("\n[ERROR] Valid ZIP file not found.")
        print("Please enter the complete ZIP path again.")

    output_root = input(
        "\nEnter OUTPUT folder path:\n"
        "(Press ENTER to create output beside the ZIP)\n> "
    ).strip().strip('"')

    if not output_root:
        output_root = os.path.join(
            os.path.dirname(os.path.abspath(zip_path)),
            "Google_Consolidated_Reports"
        )

    timestamp = datetime.now().strftime(
        "%Y%m%d_%H%M%S"
    )

    output_dir = os.path.join(
        os.path.abspath(output_root),
        f"Report_{timestamp}"
    )

    os.makedirs(output_dir, exist_ok=True)

    return (
        os.path.abspath(zip_path),
        output_dir
    )


# ============================================================
# SAFE ZIP EXTRACTION
# ============================================================

def safe_extract(zip_path, destination):
    os.makedirs(destination, exist_ok=True)

    destination_abs = os.path.abspath(destination)

    with zipfile.ZipFile(zip_path, "r") as archive:

        for member in archive.infolist():

            target = os.path.abspath(
                os.path.join(
                    destination,
                    member.filename
                )
            )

            if not (
                target == destination_abs
                or target.startswith(
                    destination_abs + os.sep
                )
            ):
                raise RuntimeError(
                    f"Unsafe ZIP entry detected: "
                    f"{member.filename}"
                )

        archive.extractall(destination)


def extract_all_nested_zips(root_dir):
    """
    Recursively extract every ZIP found under root_dir.
    """

    processed = set()
    count = 0

    while True:

        new_found = False

        for root, dirs, files in os.walk(root_dir):

            for filename in files:

                if not filename.lower().endswith(".zip"):
                    continue

                zip_path = os.path.abspath(
                    os.path.join(root, filename)
                )

                if zip_path in processed:
                    continue

                processed.add(zip_path)

                name_without_ext = os.path.splitext(
                    filename
                )[0]

                destination = os.path.join(
                    root,
                    f"__EXTRACTED__{name_without_ext}"
                )

                try:

                    safe_extract(
                        zip_path,
                        destination
                    )

                    count += 1
                    new_found = True

                    print(
                        f"[+] Nested ZIP extracted: "
                        f"{filename}"
                    )

                except Exception as exc:

                    print(
                        f"[WARNING] Could not extract "
                        f"{filename}: {exc}"
                    )

        if not new_found:
            break

    return count


# ============================================================
# BASIC HELPERS
# ============================================================

def clean_lines(text):
    return [
        line.strip()
        for line in text.splitlines()
        if line.strip()
    ]


def normalize_key(value):
    return " ".join(
        str(value)
        .strip()
        .lower()
        .split()
    )


def normalize_ip(ip):
    if ip is None:
        return ""

    value = str(ip).strip()

    # Google/export data may contain escaped IPv6 colons.
    value = value.replace("\\:", ":")
    value = value.replace("\\\\", "\\")

    return value


# ============================================================
# DATE/TIME FORMATTERS
# ============================================================

def is_special_ipv6(ip):
    ip = normalize_ip(ip).lower()

    return (
        ip.startswith("2405:")
        or ip.startswith("2409:")
    )


def is_airtel_ipv6(ip):
    """Airtel IPv6 block starts with 2401:"""
    return normalize_ip(ip).lower().startswith("2401:")


def parse_date(value):
    value = str(value or "").strip()

    if not value:
        return None

    if value.endswith(".0"):
        value = value[:-2]

    formats = [
        "%Y%m%d",
        "%d%m%Y",
        "%d/%m/%Y",
        "%Y-%m-%d",
        "%d-%m-%Y",
        "%d-%b-%Y",
        "%d-%B-%Y",
    ]

    for fmt in formats:

        try:
            return datetime.strptime(
                value,
                fmt
            )

        except ValueError:
            pass

    return None


def format_date(value, ip):
    dt = parse_date(value)

    if dt is None:
        return str(value or "").strip()

    if is_special_ipv6(ip):
        # 2405/2409 -> YYYYMMDD
        return dt.strftime("%Y%m%d")

    if is_airtel_ipv6(ip):
        # 2401 (Airtel) -> DD/MMM/YYYY
        return dt.strftime("%d/%b/%Y")

    # Other IP -> DD-MMM-YYYY
    return dt.strftime("%d-%b-%Y")


def parse_time(value):
    value = str(value or "").strip()

    if not value:
        return None

    if value.endswith(".0"):
        value = value[:-2]

    formats = [
        "%H%M%S",
        "%H:%M:%S",
        "%H%M",
        "%H:%M",
    ]

    for fmt in formats:

        try:
            return datetime.strptime(
                value,
                fmt
            )

        except ValueError:
            pass

    return None


def format_time(value, ip):
    dt = parse_time(value)

    if dt is None:
        return str(value or "").strip()

    if is_special_ipv6(ip):
        # 2405/2409 -> HHMMSS
        return dt.strftime("%H%M%S")

    # Other IP (including Airtel 2401:) -> HH:MM:SS
    return dt.strftime("%H:%M:%S")


# ============================================================
# GENERIC IP DETAIL CSV PARSER
# ============================================================

def read_csv_flexible(file_path):
    encodings = [
        "utf-8-sig",
        "utf-8",
        "cp1252",
    ]

    for encoding in encodings:

        try:

            df = pd.read_csv(
                file_path,
                dtype=str,
                keep_default_na=False,
                encoding=encoding
            )

            if len(df.columns) > 1:
                return df

        except Exception:
            pass

    try:

        return pd.read_csv(
            file_path,
            dtype=str,
            keep_default_na=False,
            sep=None,
            engine="python",
            encoding="utf-8-sig"
        )

    except Exception:
        return pd.DataFrame()


def detect_ip_detail_csv(file_path):
    df = read_csv_flexible(file_path)

    if df.empty:
        return []

    columns = {
        normalize_key(column): column
        for column in df.columns
    }

    required = [
        "ip address",
        "f date",
        "f time",
        "t date",
        "t time",
    ]

    if not all(
        item in columns
        for item in required
    ):
        return []

    type_column = columns.get("type")

    records = []

    for _, row in df.iterrows():

        ip = normalize_ip(
            row.get(
                columns["ip address"],
                ""
            )
        )

        if not ip:
            continue

        f_date = format_date(
            row.get(columns["f date"], ""),
            ip
        )

        f_time = format_time(
            row.get(columns["f time"], ""),
            ip
        )

        t_date = format_date(
            row.get(columns["t date"], ""),
            ip
        )

        t_time = format_time(
            row.get(columns["t time"], ""),
            ip
        )

        records.append({
            "TYPE": (
                row.get(type_column, "")
                if type_column
                else ""
            ),
            "IP Address": ip,
            "F DATE": f_date,
            "F TIME": f_time,
            "T DATE": t_date,
            "T TIME": t_time,
            "FROM DATE": (
                f"{f_date} {f_time}".strip()
            ),
            "TO DATE": (
                f"{t_date} {t_time}".strip()
            ),
            "Source CSV": os.path.basename(
                file_path
            ),
        })

    return records



def adjust_ip_record_times(record):
    """
    Apply the requested +/- 1 second rule.

    FROM = source timestamp - 1 second
    TO   = source timestamp + 1 second

    For explicit F/T records:
        FROM is based on F DATE + F TIME
        TO is based on T DATE + T TIME

    For fallback Google Activity records:
        both sides are based on the source timestamp.

    The date is recalculated after +/- 1 second, so midnight
    rollover is handled correctly.
    """

    ip = normalize_ip(
        record.get("IP Address", "")
    )

    special = is_special_ipv6(ip)

    f_date = str(
        record.get("F DATE", "")
    ).strip()

    f_time = str(
        record.get("F TIME", "")
    ).strip()

    t_date = str(
        record.get("T DATE", "")
    ).strip()

    t_time = str(
        record.get("T TIME", "")
    ).strip()

    # Try to build FROM/TO from the existing normalized values.
    f_dt = None
    t_dt = None

    parsed_f_date = parse_date(f_date)
    parsed_f_time = parse_time(f_time)

    if parsed_f_date and parsed_f_time:
        f_dt = datetime.combine(
            parsed_f_date.date(),
            parsed_f_time.time()
        )

    parsed_t_date = parse_date(t_date)
    parsed_t_time = parse_time(t_time)

    if parsed_t_date and parsed_t_time:
        t_dt = datetime.combine(
            parsed_t_date.date(),
            parsed_t_time.time()
        )

    # Fallback: use original source timestamp if available.
    source_dt = parse_google_timestamp(
        record.get("Timestamp UTC", "")
    )

    if source_dt is not None:
        # Convert aware datetime to naive UTC for formatting.
        if source_dt.tzinfo is not None:
            source_dt = source_dt.astimezone(
                timezone.utc
            ).replace(tzinfo=None)

        if f_dt is None:
            f_dt = source_dt

        if t_dt is None:
            t_dt = source_dt

    if f_dt is None or t_dt is None:
        return record

    f_dt = f_dt - timedelta(seconds=1)
    t_dt = t_dt + timedelta(seconds=1)

    if special:
        record["F DATE"] = f_dt.strftime(
            "%Y%m%d"
        )
        record["F TIME"] = f_dt.strftime(
            "%H%M%S"
        )
        record["T DATE"] = t_dt.strftime(
            "%Y%m%d"
        )
        record["T TIME"] = t_dt.strftime(
            "%H%M%S"
        )

    else:
        record["F DATE"] = f_dt.strftime(
            "%d-%b-%Y"
        )
        record["F TIME"] = f_dt.strftime(
            "%H:%M:%S"
        )
        record["T DATE"] = t_dt.strftime(
            "%d-%b-%Y"
        )
        record["T TIME"] = t_dt.strftime(
            "%H:%M:%S"
        )

    record["FROM DATE"] = (
        f"{record['F DATE']} {record['F TIME']}"
    )

    record["TO DATE"] = (
        f"{record['T DATE']} {record['T TIME']}"
    )

    return record


def remove_duplicate_ips(records):
    """
    Remove duplicate IPs.
    First occurrence is retained.
    """

    output = []
    seen = set()

    for record in records:

        ip = normalize_ip(
            record.get("IP Address", "")
        )

        if not ip:
            continue

        key = ip.lower()

        if key in seen:
            continue

        seen.add(key)

        record["IP Address"] = ip

        # Apply requested FROM -1 second / TO +1 second.
        record = adjust_ip_record_times(record)

        output.append(record)

    return output


# ============================================================
# GOOGLE SUBSCRIBER / DEVICE HTML PARSER
# ============================================================

def parse_html_file(
    file_path,
    subscriber,
    device_records,
    google_ip_rows
):

    try:

        with open(
            file_path,
            "r",
            encoding="utf-8",
            errors="replace"
        ) as file:

            html = file.read()

    except Exception as exc:

        print(
            f"[WARNING] HTML read failed: "
            f"{file_path} -> {exc}"
        )

        return

    soup = BeautifulSoup(
        html,
        "html.parser"
    )

    text = soup.get_text("\n")
    upper_text = text.upper()
    lines = clean_lines(text)

    # --------------------------------------------------------
    # GOOGLE SUBSCRIBER INFORMATION
    # --------------------------------------------------------

    if "GOOGLE SUBSCRIBER INFORMATION" in upper_text:

        print(
            f"[+] Subscriber HTML: "
            f"{os.path.basename(file_path)}"
        )

        subscriber_keys = [
            "Google Account ID",
            "Name",
            "Given Name",
            "Family Name",
            "e-Mail",
            "Created on",
            "Terms of Service IP",
            "Terms of Service Country",
            "Status",
            "Last Updated Date",
            "Recovery SMS",
            "User Phone Numbers",
            "Birthday (Month Day, Year)",
            "Services",
        ]

        wanted = {
            normalize_key(key): key
            for key in subscriber_keys
        }

        for line in lines:

            if ":" not in line:
                continue

            key, value = line.split(
                ":",
                1
            )

            normalized = normalize_key(key)

            if normalized in wanted:

                subscriber[
                    wanted[normalized]
                ] = value.strip()

        # ----------------------------------------------------
        # GOOGLE IP ACTIVITY TABLE
        # ----------------------------------------------------

        for table in soup.find_all("table"):

            rows = table.find_all("tr")

            if not rows:
                continue

            header_cells = rows[0].find_all(
                ["th", "td"]
            )

            headers = [
                normalize_key(
                    cell.get_text(
                        " ",
                        strip=True
                    )
                )
                for cell in header_cells
            ]

            if (
                "timestamp" not in headers
                or "ip address" not in headers
                or "activity type" not in headers
            ):
                continue

            timestamp_index = headers.index(
                "timestamp"
            )

            ip_index = headers.index(
                "ip address"
            )

            activity_index = headers.index(
                "activity type"
            )

            android_index = (
                headers.index("android id")
                if "android id" in headers
                else None
            )

            apple_index = (
                headers.index("apple ios idfv")
                if "apple ios idfv" in headers
                else None
            )

            for row in rows[1:]:

                cells = [
                    cell.get_text(
                        " ",
                        strip=True
                    )
                    for cell in row.find_all(
                        ["td", "th"]
                    )
                ]

                required_max = max(
                    timestamp_index,
                    ip_index,
                    activity_index
                )

                if len(cells) <= required_max:
                    continue

                ip = normalize_ip(
                    cells[ip_index]
                )

                if not ip:
                    continue

                google_ip_rows.append({
                    "Timestamp UTC": (
                        cells[timestamp_index]
                    ),
                    "IP Address": ip,
                    "Activity Type": (
                        cells[activity_index]
                    ),
                    "Android ID": (
                        cells[android_index]
                        if android_index is not None
                        and android_index < len(cells)
                        else ""
                    ),
                    "Apple iOS IDFV": (
                        cells[apple_index]
                        if apple_index is not None
                        and apple_index < len(cells)
                        else ""
                    ),
                    "Source HTML": os.path.basename(
                        file_path
                    ),
                })

    # --------------------------------------------------------
    # ANDROID DEVICE CONFIGURATION
    # --------------------------------------------------------

    elif (
        "ANDROID DEVICE CONFIGURATION SERVICE DATA"
        in upper_text
        or "ANDROID DEVICE CONFIGURATION"
        in upper_text
    ):

        print(
            f"[+] Device HTML: "
            f"{os.path.basename(file_path)}"
        )

        device_keys = [
            "Android ID",
            "IMEI(s)",
            "MEID(s)",
            "Serial Number(s)",
            "Users",
            "Locale",
            "Timezone",
            "Model",
            "Brand",
            "Manufacturer",
            "Device Type",
            "Device Sub-type",
            "Device",
            "Product",
            "Partner Client ID",
            "Time of Last Data Connection",
            "IP address from Last Data Connection",
            "Reason for Last Data Connection",
            "Last Carrier ID",
        ]

        wanted = {
            normalize_key(key): key
            for key in device_keys
        }

        device = {}

        for line in lines:

            if ":" not in line:
                continue

            key, value = line.split(
                ":",
                1
            )

            normalized = normalize_key(key)

            if normalized in wanted:

                device[
                    wanted[normalized]
                ] = value.strip()

        if device:

            device["Source HTML"] = (
                os.path.basename(file_path)
            )

            device_records.append(
                device
            )



# ============================================================
# BUILD IP DETAILS FROM GOOGLE IP ACTIVITY
# ============================================================

def parse_google_timestamp(value):
    """
    Parse Google timestamp such as:
        2026-08-13 14:43:05 Z
        2026-08-13T14:43:05Z
        2026-08-13 14:43:05+00:00
    """
    value = str(value or "").strip()

    if not value:
        return None

    value = value.replace(" UTC", "+00:00")
    value = value.replace(" Z", "+00:00")
    value = value.replace("Z", "+00:00")

    try:
        dt = datetime.fromisoformat(value)

        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)

        return dt

    except Exception:
        pass

    for fmt in [
        "%Y-%m-%d %H:%M:%S",
        "%Y-%m-%dT%H:%M:%S",
        "%Y-%m-%d %H:%M:%S Z",
        "%Y-%m-%dT%H:%M:%SZ",
    ]:
        try:
            dt = datetime.strptime(value, fmt)
            return dt.replace(tzinfo=timezone.utc)
        except ValueError:
            pass

    return None


def build_ip_details_from_google_activity(
    google_ip_rows
):
    """
    Fallback for Google Subscriber exports.

    The supplied Google ZIP may contain only:
        Timestamp
        IP Address
        Activity Type
        Android ID
        Apple iOS IDFV

    It does NOT contain F DATE/F TIME/T DATE/T TIME.

    In that situation, create one IP-detail record per
    unique IP using the Google activity timestamp as both
    FROM and TO.

    Formatting:
        2405:/2409: -> YYYYMMDD HHMMSS
        Other IPs   -> DD-MMM-YYYY HH:MM:SS
    """

    records = []

    for row in google_ip_rows:

        ip = normalize_ip(
            row.get("IP Address", "")
        )

        if not ip:
            continue

        dt = parse_google_timestamp(
            row.get("Timestamp UTC", "")
        )

        if dt is None:
            continue

        if is_special_ipv6(ip):

            date_part = dt.strftime(
                "%Y%m%d"
            )

            time_part = dt.strftime(
                "%H%M%S"
            )

        elif is_airtel_ipv6(ip):

            date_part = dt.strftime(
                "%d/%b/%Y"
            )

            time_part = dt.strftime(
                "%H:%M:%S"
            )

        else:

            date_part = dt.strftime(
                "%d-%b-%Y"
            )

            time_part = dt.strftime(
                "%H:%M:%S"
            )

        records.append({
            "TYPE": "IPV6"
            if ":" in ip
            else "IPV4",
            "IP Address": ip,
            "F DATE": date_part,
            "F TIME": time_part,
            "T DATE": date_part,
            "T TIME": time_part,
            "FROM DATE": (
                f"{date_part} {time_part}"
            ),
            "TO DATE": (
                f"{date_part} {time_part}"
            ),
            "Source": "Google IP Activity",
            "Activity Type": row.get(
                "Activity Type",
                ""
            ),
        })

    return records


def extract_ip_detail_tables_from_html(
    file_path
):
    """
    Extra fallback:
    Detect an HTML table containing:
        IP Address
        F DATE
        F TIME
        T DATE
        T TIME
    """

    records = []

    try:
        with open(
            file_path,
            "r",
            encoding="utf-8",
            errors="replace"
        ) as file:
            html = file.read()

    except Exception:
        return records

    soup = BeautifulSoup(
        html,
        "html.parser"
    )

    for table in soup.find_all("table"):

        rows = table.find_all("tr")

        if not rows:
            continue

        headers = [
            normalize_key(
                cell.get_text(
                    " ",
                    strip=True
                )
            )
            for cell in rows[0].find_all(
                ["th", "td"]
            )
        ]

        required = {
            "ip address",
            "f date",
            "f time",
            "t date",
            "t time",
        }

        if not required.issubset(
            set(headers)
        ):
            continue

        index = {
            header: headers.index(header)
            for header in required
        }

        type_index = (
            headers.index("type")
            if "type" in headers
            else None
        )

        for row in rows[1:]:

            cells = [
                cell.get_text(
                    " ",
                    strip=True
                )
                for cell in row.find_all(
                    ["td", "th"]
                )
            ]

            if not cells:
                continue

            try:
                ip = normalize_ip(
                    cells[index["ip address"]]
                )

                if not ip:
                    continue

                records.append({
                    "TYPE": (
                        cells[type_index]
                        if type_index is not None
                        and type_index < len(cells)
                        else (
                            "IPV6"
                            if ":" in ip
                            else "IPV4"
                        )
                    ),
                    "IP Address": ip,
                    "F DATE": format_date(
                        cells[index["f date"]],
                        ip
                    ),
                    "F TIME": format_time(
                        cells[index["f time"]],
                        ip
                    ),
                    "T DATE": format_date(
                        cells[index["t date"]],
                        ip
                    ),
                    "T TIME": format_time(
                        cells[index["t time"]],
                        ip
                    ),
                    "FROM DATE": "",
                    "TO DATE": "",
                    "Source": os.path.basename(
                        file_path
                    ),
                })

                records[-1]["FROM DATE"] = (
                    f"{records[-1]['F DATE']} "
                    f"{records[-1]['F TIME']}"
                )

                records[-1]["TO DATE"] = (
                    f"{records[-1]['T DATE']} "
                    f"{records[-1]['T TIME']}"
                )

            except (IndexError, ValueError):
                continue

    return records


# ============================================================
# MAIN DATA COLLECTION
# ============================================================

def collect_data(extraction_root):

    subscriber = {}

    device_records = []

    google_ip_rows = []

    target_records = []

    formatted_ip_records = []

    html_ip_records = []

    html_count = 0
    csv_count = 0

    for root, dirs, files in os.walk(
        extraction_root
    ):

        for filename in files:

            file_path = os.path.join(
                root,
                filename
            )

            lower = filename.lower()

            if lower.endswith(
                (".html", ".htm")
            ):

                html_count += 1

                parse_html_file(
                    file_path,
                    subscriber,
                    device_records,
                    google_ip_rows
                )

                html_ip_records.extend(
                    extract_ip_detail_tables_from_html(
                        file_path
                    )
                )

            elif lower.endswith(".csv"):

                csv_count += 1

                # IP detail CSV
                ip_records = (
                    detect_ip_detail_csv(
                        file_path
                    )
                )

                if ip_records:

                    print(
                        f"[+] IP detail CSV: "
                        f"{filename} -> "
                        f"{len(ip_records)} records"
                    )

                    formatted_ip_records.extend(
                        ip_records
                    )

                # TargetAsPhone CSV
                if (
                    "targetasphone"
                    in lower
                ):

                    df = read_csv_flexible(
                        file_path
                    )

                    if not df.empty:

                        print(
                            f"[+] TargetAsPhone CSV: "
                            f"{filename} -> "
                            f"{len(df)} records"
                        )

                        for record in (
                            df.to_dict(
                                "records"
                            )
                        ):

                            record[
                                "Source CSV"
                            ] = filename

                            target_records.append(
                                record
                            )

    # If an explicit F DATE/F TIME/T DATE/T TIME
    # source was found, use it.
    #
    # If not, use IP detail tables from HTML.
    if not formatted_ip_records and html_ip_records:
        formatted_ip_records.extend(
            html_ip_records
        )

    # Final fallback for the Google ZIP structure used
    # in this investigation: SubscriberInfo.html contains
    # an IP ACTIVITY table but no F/T columns.
    if not formatted_ip_records:
        formatted_ip_records.extend(
            build_ip_details_from_google_activity(
                google_ip_rows
            )
        )

    # Remove duplicate IPs globally.
    formatted_ip_records = (
        remove_duplicate_ips(
            formatted_ip_records
        )
    )

    return {
        "subscriber": subscriber,
        "devices": device_records,
        "google_ip": google_ip_rows,
        "target_phone": target_records,
        "ip_details": formatted_ip_records,
        "html_count": html_count,
        "csv_count": csv_count,
    }


# ============================================================
# GOOGLE IP UTC -> IST
# ============================================================

def convert_google_ip_to_ist(rows):

    ist = timezone(
        timedelta(hours=5, minutes=30)
    )

    for row in rows:

        raw = str(
            row.get(
                "Timestamp UTC",
                ""
            )
        ).strip()

        row["Timestamp IST"] = ""

        if not raw:
            continue

        value = raw

        value = value.replace(
            " UTC",
            "+00:00"
        )

        value = value.replace(
            " Z",
            "+00:00"
        )

        try:

            dt = datetime.fromisoformat(
                value
            )

            if dt.tzinfo is None:

                dt = dt.replace(
                    tzinfo=timezone.utc
                )

            dt = dt.astimezone(ist)

            row["Timestamp IST"] = (
                dt.strftime(
                    "%d-%m-%Y %H:%M:%S"
                )
            )

        except Exception:

            for fmt in [
                "%Y-%m-%d %H:%M:%S",
                "%Y-%m-%d %H:%M:%S UTC",
                "%Y-%m-%dT%H:%M:%S",
            ]:

                try:

                    dt = datetime.strptime(
                        raw,
                        fmt
                    )

                    dt = dt.replace(
                        tzinfo=timezone.utc
                    )

                    dt = dt.astimezone(ist)

                    row["Timestamp IST"] = (
                        dt.strftime(
                            "%d-%m-%Y %H:%M:%S"
                        )
                    )

                    break

                except ValueError:
                    pass

    return rows


# ============================================================
# EXCEL
# ============================================================

def format_excel_sheet(sheet):

    if sheet.max_row < 1:
        return

    for cell in sheet[1]:

        cell.font = Font(
            bold=True,
            color="FFFFFF"
        )

        cell.fill = PatternFill(
            "solid",
            fgColor="4472C4"
        )

        cell.alignment = Alignment(
            horizontal="center",
            vertical="center"
        )

    sheet.freeze_panes = "A2"

    if sheet.max_row > 1:
        sheet.auto_filter.ref = (
            sheet.dimensions
        )

    for column in range(
        1,
        sheet.max_column + 1
    ):

        values = []

        for row in range(
            1,
            min(
                sheet.max_row,
                500
            ) + 1
        ):

            values.append(
                str(
                    sheet.cell(
                        row,
                        column
                    ).value
                    or ""
                )
            )

        width = min(
            max(
                [len(value) for value in values]
                + [10]
            ) + 2,
            70
        )

        sheet.column_dimensions[
            get_column_letter(column)
        ].width = width


def add_dataframe_sheet(
    workbook,
    name,
    dataframe
):

    sheet = workbook.create_sheet(
        name
    )

    if dataframe.empty:

        sheet.append([
            "No records found"
        ])

        return

    sheet.append(
        list(dataframe.columns)
    )

    for row in dataframe.itertuples(
        index=False,
        name=None
    ):

        sheet.append(
            list(row)
        )

    format_excel_sheet(
        sheet
    )


def create_excel_report(
    output_dir,
    zip_path,
    data,
    ip_info=None,
):

    output_file = os.path.join(
        output_dir,
        "Consolidated_Google_Investigation_Report.xlsx"
    )

    subscriber = data["subscriber"]

    df_google_ip = pd.DataFrame(
        data["google_ip"]
    )

    df_devices = pd.DataFrame(
        data["devices"]
    )

    df_target = pd.DataFrame(
        data["target_phone"]
    )

    df_ip_details = pd.DataFrame(
        data["ip_details"]
    )

    workbook = Workbook()

    summary = workbook.active
    summary.title = "Executive Summary"

    total_ip = len(df_google_ip)

    successful = 0
    failed = 0
    unique_google_ips = 0

    if (
        not df_google_ip.empty
        and "Activity Type"
        in df_google_ip.columns
    ):

        activity = (
            df_google_ip[
                "Activity Type"
            ]
            .astype(str)
            .str.strip()
            .str.lower()
        )

        successful = int(
            activity.isin([
                "login",
                "successful login",
                "success",
            ]).sum()
        )

        failed = int(
            activity.isin([
                "failed login",
                "failed",
                "failure",
            ]).sum()
        )

        unique_google_ips = int(
            df_google_ip[
                "IP Address"
            ]
            .astype(str)
            .nunique()
        )

    connected = 0

    if (
        not df_target.empty
        and "Connection Status"
        in df_target.columns
    ):

        connected = int(
            df_target[
                "Connection Status"
            ]
            .astype(str)
            .str.upper()
            .eq("CONNECTED")
            .sum()
        )

    summary_rows = [
        [
            "CONSOLIDATED GOOGLE "
            "ACCOUNT / DEVICE REPORT",
            ""
        ],
        [
            "Source ZIP",
            os.path.basename(zip_path)
        ],
        [
            "Primary Google Account",
            subscriber.get(
                "e-Mail",
                ""
            )
        ],
        [
            "Google Account ID",
            subscriber.get(
                "Google Account ID",
                ""
            )
        ],
        [
            "Account Name",
            subscriber.get(
                "Name",
                ""
            )
        ],
        [
            "Account Status",
            subscriber.get(
                "Status",
                ""
            )
        ],
        [
            "Created On UTC",
            subscriber.get(
                "Created on",
                ""
            )
        ],
        [
            "IP Activity Records",
            total_ip
        ],
        [
            "Successful Login Records",
            successful
        ],
        [
            "Failed Login Records",
            failed
        ],
        [
            "Unique Google IPs",
            unique_google_ips
        ],
        [
            "Formatted IP Detail Records",
            len(df_ip_details)
        ],
        [
            "Unique Formatted IPs",
            len(df_ip_details)
        ],
        [
            "Phone Association Records",
            len(df_target)
        ],
        [
            "Connected Phone Associations",
            connected
        ],
        [
            "Device Records",
            len(df_devices)
        ],
    ]

    for row in summary_rows:
        summary.append(row)

    for cell in summary[1]:

        cell.fill = PatternFill(
            "solid",
            fgColor="1F4E78"
        )

        cell.font = Font(
            color="FFFFFF",
            bold=True,
            size=14
        )

    summary.column_dimensions["A"].width = 42
    summary.column_dimensions["B"].width = 75

    # Subscriber details
    subscriber_rows = []

    for key in [
        "Google Account ID",
        "Name",
        "Given Name",
        "Family Name",
        "e-Mail",
        "Created on",
        "Terms of Service IP",
        "Terms of Service Country",
        "Status",
        "Last Updated Date",
        "Recovery SMS",
        "User Phone Numbers",
        "Birthday (Month Day, Year)",
        "Services",
    ]:

        if key in subscriber:

            subscriber_rows.append({
                "Field": key,
                "Value": subscriber[key],
            })

    add_dataframe_sheet(
        workbook,
        "Subscriber",
        pd.DataFrame(
            subscriber_rows,
            columns=["Field", "Value"]
        )
    )

    add_dataframe_sheet(
        workbook,
        "IP Activity",
        df_google_ip
    )

    add_dataframe_sheet(
        workbook,
        "IP Details",
        df_ip_details
    )

    add_dataframe_sheet(
        workbook,
        "Phone Associations",
        df_target
    )

    add_dataframe_sheet(
        workbook,
        "Devices",
        df_devices
    )

    if ip_info:
        geo_rows = []
        for ip in sorted(ip_info.keys()):
            info = ip_info.get(ip) or {}
            geo_rows.append({
                "IP Address": ip,
                "City": info.get("city", ""),
                "Region": info.get("region", ""),
                "Country": info.get("country", ""),
                "Postal": info.get("postal", ""),
                "Coordinates": info.get("loc", ""),
                "ISP / Organization": info.get("org", ""),
                "Hostname": info.get("hostname", ""),
                "Timezone": info.get("timezone", ""),
                "Notes": (
                    info.get("note", "")
                    or info.get("error", "")
                    or ("bogon / private" if info.get("bogon") else "")
                ),
            })
        add_dataframe_sheet(
            workbook,
            "IP Geolocation",
            pd.DataFrame(geo_rows),
        )

    workbook.save(
        output_file
    )

    return output_file


# ============================================================
# DOCX TABLE
# ============================================================

def add_docx_table(
    document,
    headers,
    rows
):

    table = document.add_table(
        rows=1,
        cols=len(headers)
    )

    table.style = "Table Grid"

    table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    for index, header in enumerate(
        headers
    ):

        cell = table.rows[0].cells[
            index
        ]

        cell.text = str(header)

        cell.vertical_alignment = (
            WD_CELL_VERTICAL_ALIGNMENT.CENTER
        )

        for run in (
            cell.paragraphs[0].runs
        ):

            run.bold = True

    for row in rows:

        cells = table.add_row().cells

        for index, value in enumerate(
            row
        ):

            cells[index].text = str(
                value
            )

    return table


# ============================================================
# DOCX REPORT
# ============================================================

def create_docx_report(
    output_dir,
    zip_path,
    data
):

    output_file = os.path.join(
        output_dir,
        "Consolidated_Google_Investigation_Report.docx"
    )

    subscriber = data["subscriber"]

    df_google_ip = pd.DataFrame(
        data["google_ip"]
    )

    df_devices = pd.DataFrame(
        data["devices"]
    )

    df_target = pd.DataFrame(
        data["target_phone"]
    )

    df_ip_details = pd.DataFrame(
        data["ip_details"]
    )

    document = Document()

    section = document.sections[0]

    section.top_margin = Inches(
        0.65
    )

    section.bottom_margin = Inches(
        0.65
    )

    section.left_margin = Inches(
        0.55
    )

    section.right_margin = Inches(
        0.55
    )

    # --------------------------------------------------------
    # TITLE
    # --------------------------------------------------------

    paragraph = document.add_paragraph()

    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    run = paragraph.add_run(
        "CONSOLIDATED GOOGLE ACCOUNT "
        "& DEVICE ANALYSIS REPORT"
    )

    run.bold = True
    run.font.size = Pt(16)

    paragraph = document.add_paragraph()

    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    run = paragraph.add_run(
        "Subscriber Information | Device Data | "
        "Phone Association | IP Analysis"
    )

    run.italic = True
    run.font.size = Pt(9)

    # --------------------------------------------------------
    # SUMMARY
    # --------------------------------------------------------

    document.add_heading(
        "1. Executive Summary",
        level=1
    )

    document.add_paragraph(
        f"The supplied Google export package contains "
        f"{len(df_google_ip)} Google IP activity records, "
        f"{len(df_devices)} device record(s), "
        f"{len(df_target)} phone association record(s), "
        f"and {len(df_ip_details)} unique formatted IP "
        f"detail record(s)."
    )

    document.add_paragraph(
        f"Primary Google account: "
        f"{subscriber.get('e-Mail', '')}."
    )

    # --------------------------------------------------------
    # SUBSCRIBER
    # --------------------------------------------------------

    document.add_heading(
        "2. Google Subscriber Details",
        level=1
    )

    subscriber_rows = []

    for key in [
        "Google Account ID",
        "Name",
        "Given Name",
        "Family Name",
        "e-Mail",
        "Created on",
        "Terms of Service IP",
        "Terms of Service Country",
        "Status",
        "Last Updated Date",
        "Recovery SMS",
        "User Phone Numbers",
        "Birthday (Month Day, Year)",
        "Services",
    ]:

        if key in subscriber:

            subscriber_rows.append([
                key,
                subscriber[key]
            ])

    add_docx_table(
        document,
        ["Field", "Value"],
        subscriber_rows
    )

    # --------------------------------------------------------
    # DEVICES
    # --------------------------------------------------------

    document.add_heading(
        "3. Android Device Correlation",
        level=1
    )

    if not df_devices.empty:

        columns = [
            "Android ID",
            "Model",
            "Brand",
            "IMEI(s)",
            "Serial Number(s)",
            "Users",
            "Timezone",
            "Time of Last Data Connection",
            "IP address from Last Data Connection",
        ]

        columns = [
            c for c in columns
            if c in df_devices.columns
        ]

        rows = []

        for _, record in (
            df_devices.iterrows()
        ):

            rows.append([
                record.get(
                    column,
                    ""
                )
                for column in columns
            ])

        add_docx_table(
            document,
            columns,
            rows
        )

    else:

        document.add_paragraph(
            "No Android device records found."
        )

    # --------------------------------------------------------
    # PHONE ASSOCIATIONS
    # --------------------------------------------------------

    document.add_heading(
        "4. Phone Number / Google Account Association",
        level=1
    )

    if not df_target.empty:

        columns = list(
            df_target.columns
        )

        rows = []

        for _, record in (
            df_target.iterrows()
        ):

            rows.append([
                record.get(
                    column,
                    ""
                )
                for column in columns
            ])

        add_docx_table(
            document,
            columns,
            rows
        )

    else:

        document.add_paragraph(
            "No TargetAsPhone association records found."
        )

    # --------------------------------------------------------
    # IP DETAILS - REQUIRED FORMAT
    # --------------------------------------------------------

    document.add_heading(
        "5. IP Details",
        level=1
    )

    document.add_paragraph(
        "IP formatting rule: 2405: and 2409: IPv6 "
        "addresses use YYYYMMDD HHMMSS; 2401: (Airtel) "
        "IPv6 addresses use DD/MMM/YYYY HH:MM:SS; all "
        "other IP addresses use DD-MMM-YYYY HH:MM:SS. "
        "Duplicate IP addresses are removed and the first "
        "occurrence is retained."
    )

    if not df_ip_details.empty:

        rows = []

        for _, record in (
            df_ip_details.iterrows()
        ):

            rows.append([
                record.get(
                    "IP Address",
                    ""
                ),
                record.get(
                    "FROM DATE",
                    ""
                ),
                record.get(
                    "TO DATE",
                    ""
                ),
            ])

        add_docx_table(
            document,
            [
                "IP Address",
                "FROM DATE",
                "TO DATE",
            ],
            rows[:100]
        )

        if len(rows) > 100:

            document.add_paragraph(
                f"First 100 IP detail records are shown "
                f"in this DOCX. The complete "
                f"{len(rows)} unique records are available "
                f"in the Excel IP Details sheet and CSV."
            )

    else:

        document.add_paragraph(
            "No IP detail CSV containing F DATE/F TIME/"
            "T DATE/T TIME columns was found."
        )

    # --------------------------------------------------------
    # GOOGLE IP ACTIVITY
    # --------------------------------------------------------

    document.add_heading(
        "6. Google IP Activity",
        level=1
    )

    if not df_google_ip.empty:

        columns = [
            c for c in [
                "Timestamp UTC",
                "Timestamp IST",
                "IP Address",
                "Activity Type",
                "Android ID",
                "Apple iOS IDFV",
            ]
            if c in df_google_ip.columns
        ]

        rows = []

        for _, record in (
            df_google_ip.head(50).iterrows()
        ):

            rows.append([
                record.get(
                    column,
                    ""
                )
                for column in columns
            ])

        add_docx_table(
            document,
            columns,
            rows
        )

        if len(df_google_ip) > 50:

            document.add_paragraph(
                "First 50 Google IP activity records are "
                "shown in this DOCX. Complete records are "
                "available in the Excel IP Activity sheet."
            )

    else:

        document.add_paragraph(
            "No Google IP activity records found."
        )

    # --------------------------------------------------------
    # INVESTIGATIVE LEADS
    # --------------------------------------------------------

    document.add_heading(
        "7. Investigative Leads / Further Verification",
        level=1
    )

    leads = [
        "Obtain ISP subscriber/KYC and CGNAT/session "
        "records for relevant IP addresses using exact "
        "timestamps.",

        "Correlate Android ID, IMEI, serial number and "
        "last data connection information with CAF, "
        "CDR, IPDR and available forensic evidence.",

        "Verify each Google account associated with the "
        "queried phone number separately.",

        "Preserve the original ZIP and calculate a "
        "cryptographic hash for evidence integrity.",

        "Review failed-login activity separately for "
        "investigative relevance.",

        "Use the Excel IP Details sheet and CSV for the "
        "complete deduplicated IP data.",
    ]

    for lead in leads:

        document.add_paragraph(
            lead,
            style="List Number"
        )

    # --------------------------------------------------------
    # SCOPE
    # --------------------------------------------------------

    document.add_heading(
        "8. Source / Scope Note",
        level=1
    )

    document.add_paragraph(
        "This report is a consolidation of the records "
        "contained in the supplied Google export ZIP. "
        "No external IP geolocation or attribution database "
        "has been used."
    )

    document.add_paragraph(
        "The Excel report contains the complete extracted "
        "data. The DOCX is a consolidated readable report. "
        "The IP_Details_Formatted.csv contains the "
        "deduplicated IP details in the requested format."
    )

    paragraph = document.add_paragraph()

    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    run = paragraph.add_run(
        "END OF REPORT"
    )

    run.bold = True

    document.save(
        output_file
    )

    return output_file


# ============================================================
# SEPARATE WORD FILES FOR IP FORMATTING
# ============================================================

def create_ip_word_file(
    output_file,
    title,
    ip_records,
    special=False
):
    """
    Create the requested Word table.

    Special 2405/2409 file:
        IPV
        IP Address
        FROM DATE
        FROM TIME
        TO DATE
        TO TIME

    Remaining IP file:
        IP Address
        FROM DATE
        DD-MMM-YYYY HH:MM:SS
        TO DATE
        DD-MMM-YYYY HH:MM:SS

    The latter uses two-line headers to match the requested
    table format.
    """

    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(14)

    table = document.add_table(
        rows=1,
        cols=6 if special else 3
    )

    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    if special:
        headers = [
            "IPV",
            "IP Address",
            "FROM DATE",
            "FROM TIME",
            "TO DATE",
            "TO TIME",
        ]
    else:
        headers = [
            "IP Address",
            "FROM DATE\nDD-MMM-YYYY or DD/MMM/YYYY HH:MM:SS",
            "TO DATE\nDD-MMM-YYYY or DD/MMM/YYYY HH:MM:SS",
        ]

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        cell.vertical_alignment = (
            WD_CELL_VERTICAL_ALIGNMENT.CENTER
        )

        for run in cell.paragraphs[0].runs:
            run.bold = True

    for record in ip_records:
        if special:
            values = [
                record.get("TYPE", "IPV6") or "IPV6",
                record.get("IP Address", ""),
                record.get("F DATE", ""),
                record.get("F TIME", ""),
                record.get("T DATE", ""),
                record.get("T TIME", ""),
            ]
        else:
            values = [
                record.get("IP Address", ""),
                record.get("FROM DATE", ""),
                record.get("TO DATE", ""),
            ]

        cells = table.add_row().cells

        for index, value in enumerate(values):
            cells[index].text = str(value)
            cells[index].vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.CENTER
            )

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(
        f"Total Unique IPs: {len(ip_records)}"
    )
    run.bold = True

    document.save(output_file)

    return output_file



def create_separate_ip_word_files(
    output_dir,
    ip_records
):
    """
    Create two separate Word files:

    1. 2405_2409_IP_Details.docx
       Contains only IPs beginning with 2405: or 2409:

    2. Remaining_IP_Details.docx
       Contains every other IP.

    Duplicate IPs have already been removed globally.
    """

    special_records = []
    remaining_records = []

    for record in ip_records:

        ip = normalize_ip(
            record.get("IP Address", "")
        ).lower()

        if (
            ip.startswith("2405:")
            or ip.startswith("2409:")
        ):
            special_records.append(record)
        else:
            remaining_records.append(record)

    special_file = os.path.join(
        output_dir,
        "2405_2409_IP_Details.docx"
    )

    remaining_file = os.path.join(
        output_dir,
        "Remaining_IP_Details.docx"
    )

    create_ip_word_file(
        special_file,
        "2405 / 2409 IPv6 IP DETAILS",
        special_records,
        special=True
    )

    create_ip_word_file(
        remaining_file,
        "REMAINING IP DETAILS",
        remaining_records,
        special=False
    )

    return (
        special_file,
        remaining_file,
        len(special_records),
        len(remaining_records),
    )


# ============================================================
# TEXT FILE OUTPUT
# ============================================================

def create_ip_text_files(
    output_dir,
    ip_records
):
    """
    Generate separate TXT files containing only IP address
    records in the same requested formatting.

    2405/2409:
        IP Address
        FROM DATE
        FROM TIME
        TO DATE
        TO TIME
        (2401: Airtel records use DD/MMM/YYYY HH:MM:SS.)

    Remaining:
        IP Address
        FROM DATE DD-MMM-YYYY HH:MM:SS
                  (2401: Airtel -> DD/MMM/YYYY HH:MM:SS)
        TO DATE   DD-MMM-YYYY HH:MM:SS
                  (2401: Airtel -> DD/MMM/YYYY HH:MM:SS)

    Excel paste compatibility:
        F TIME / T TIME values that look like numbers (e.g.
        0042405) are prefixed with a single apostrophe so
        Excel pastes them as TEXT, preserving leading zeros.
        The apostrophe is invisible in Excel cell display.

    Duplicate IPs are already removed before this function.
    """

    def _excel_text(value):
        """
        Prefix a numeric-looking time value with an apostrophe
        so that pasting it into Excel preserves leading zeros
        (e.g. 0042405 stays 0042405 instead of becoming 42405).
        Empty values pass through unchanged.
        """
        s = str(value or "").strip()
        if not s:
            return ""
        if s.isdigit():
            return "'" + s
        return s

    special_records = []
    remaining_records = []

    for record in ip_records:

        ip = normalize_ip(
            record.get("IP Address", "")
        ).lower()

        if (
            ip.startswith("2405:")
            or ip.startswith("2409:")
        ):
            special_records.append(record)
        else:
            remaining_records.append(record)

    special_file = os.path.join(
        output_dir,
        "2405_2409_IP_Details.txt"
    )

    remaining_file = os.path.join(
        output_dir,
        "Remaining_IP_Details.txt"
    )

    # --------------------------------------------------------
    # 2405 / 2409 TXT
    # --------------------------------------------------------

    with open(
        special_file,
        "w",
        encoding="utf-8"
    ) as file:

        file.write(
            "# Format: tab-separated. F TIME / T TIME columns "
            "may carry a leading apostrophe (e.g. '0042405) "
            "so leading zeros survive an Excel paste. The "
            "apostrophe is hidden in Excel cell display.\n"
        )
        file.write(
            "IPV\tIP Address\tFROM DATE\tFROM TIME\t"
            "TO DATE\tTO TIME\n"
        )

        for record in special_records:

            file.write(
                f"{record.get('TYPE', 'IPV6') or 'IPV6'}\t"
                f"{record.get('IP Address', '')}\t"
                f"{record.get('F DATE', '')}\t"
                f"{_excel_text(record.get('F TIME', ''))}\t"
                f"{record.get('T DATE', '')}\t"
                f"{_excel_text(record.get('T TIME', ''))}\n"
            )

    # --------------------------------------------------------
    # Remaining IP TXT
    # --------------------------------------------------------

    with open(
        remaining_file,
        "w",
        encoding="utf-8"
    ) as file:

        file.write(
            "IP Address\tFROM DATE DD-MMM-YYYY HH:MM:SS\t"
            "TO DATE DD-MMM-YYYY HH:MM:SS\n"
        )

        for record in remaining_records:

            file.write(
                f"{record.get('IP Address', '')}\t"
                f"{record.get('FROM DATE', '')}\t"
                f"{record.get('TO DATE', '')}\n"
            )

    return (
        special_file,
        remaining_file,
        len(special_records),
        len(remaining_records),
    )


# ============================================================
# IP GEOLOCATION LOOKUP (ipinfo.io)
# ============================================================

def lookup_ip_info(
    unique_ips,
    token=None,
    max_per_run=500,
    sleep_seconds=0.06,
):
    """
    Look up geolocation info for each IP using ipinfo.io.

    Args:
        unique_ips: iterable of IP address strings.
        token: ipinfo.io API token. If None, looked up from
            the IPINFO_TOKEN environment variable.
        max_per_run: cap on lookups per pipeline run (safety).
        sleep_seconds: pause between requests to respect rate
            limits (free tier: 50k/month with token).

    Returns:
        dict mapping IP -> info dict. Failed lookups return
        {"ip": ip, "error": "..."}.

    Requires:
        Internet access to api.ipinfo.io. If no token is
        available the function still runs but hits the
        unauthenticated limit (~1k/day).
    """
    if not unique_ips:
        return {}

    if not token:
        token = os.environ.get("IPINFO_TOKEN")

    result = {}

    try:
        import urllib.request
        import urllib.error
        import json as _json
    except ImportError:
        # stdlib modules should always be present; this is just
        # defensive in case of unusual Python builds.
        return result

    looked = 0
    for ip in unique_ips:
        if looked >= max_per_run:
            break

        ip = str(ip or "").strip()
        if not ip:
            continue

        looked += 1

        url = "https://ipinfo.io/" + ip + "/json"
        if token:
            sep = "&" if "?" in url else "?"
            url = url + sep + "token=" + token

        try:
            req = urllib.request.Request(
                url,
                headers={"User-Agent": "IRMS/1.0"},
            )
            with urllib.request.urlopen(req, timeout=12) as resp:
                data = _json.loads(resp.read().decode("utf-8"))
            result[ip] = data
        except Exception as exc:
            result[ip] = {"ip": ip, "error": str(exc)}

        if sleep_seconds:
            import time as _time
            _time.sleep(sleep_seconds)

    return result


def create_ip_geolocation_excel(
    output_dir,
    ip_info,
):
    """
    Build a single-sheet Excel workbook with one row per
    looked-up IP. Columns cover city / region / country /
    postal / coordinates / ISP / hostname / timezone / notes.
    """
    output_file = os.path.join(
        output_dir,
        "IP_Geolocation_Lookup.xlsx",
    )

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "IP Geolocation"

    headers = [
        "IP Address",
        "City",
        "Region",
        "Country",
        "Postal",
        "Coordinates (Lat,Long)",
        "ISP / Organization",
        "Hostname",
        "Timezone",
        "Notes",
    ]
    sheet.append(headers)

    for ip in sorted(ip_info.keys()):
        info = ip_info.get(ip) or {}
        sheet.append([
            ip,
            info.get("city", "") or "",
            info.get("region", "") or "",
            info.get("country", "") or "",
            info.get("postal", "") or "",
            info.get("loc", "") or "",
            info.get("org", "") or "",
            info.get("hostname", "") or "",
            info.get("timezone", "") or "",
            (
                info.get("note", "")
                or info.get("error", "")
                or ("bogon / private" if info.get("bogon") else "")
                or ""
            ),
        ])

    format_excel_sheet(sheet)
    workbook.save(output_file)
    return output_file


# ============================================================
# CSV OUTPUT
# ============================================================

def create_ip_csv(
    output_dir,
    ip_records
):

    output_file = os.path.join(
        output_dir,
        "IP_Details_Formatted.csv"
    )

    columns = [
        "IP Address",
        "FROM DATE",
        "TO DATE",
    ]

    df = pd.DataFrame(
        ip_records,
        columns=columns
    )

    df.to_csv(
        output_file,
        index=False,
        encoding="utf-8-sig"
    )

    return output_file


# ============================================================
# PROGRAMMATIC PIPELINE (used by the web UI)
# Keeps main() intact.
# ============================================================

class PipelineLogger:
    """Captures print() output into a list and forwards lines to a callback."""

    def __init__(self, callback=None):
        self._buffer = []
        self._callback = callback

    def write(self, text):
        text = str(text)
        if not text:
            return
        for line in text.splitlines():
            line = line.rstrip()
            if not line:
                continue
            self._buffer.append(line)
            if self._callback:
                try:
                    self._callback(line)
                except Exception:
                    pass

    def flush(self):
        pass

    def lines(self):
        return list(self._buffer)


def run_pipeline(
    zip_path,
    output_dir,
    progress_callback=None,
    max_zip_size_bytes=2 * 1024 * 1024 * 1024,
    ipinfo_token=None,
):
    """
    Programmatic version of main().

    Accepts zip_path and output_dir directly, streams progress
    lines via progress_callback(line), and returns a dict
    containing all generated output files plus stats.

    The optional ipinfo_token enables IP geolocation lookups
    via the ipinfo.io API. When provided (or set in the
    IPINFO_TOKEN environment variable) the pipeline produces an
    extra IP_Geolocation_Lookup.xlsx file.

    Raises on failure so the caller can surface the error.
    """
    if not os.path.isfile(zip_path):
        raise FileNotFoundError(f"ZIP file not found: {zip_path}")

    if not zipfile.is_zipfile(zip_path):
        raise zipfile.BadZipFile(f"Not a valid ZIP file: {zip_path}")

    if os.path.getsize(zip_path) > max_zip_size_bytes:
        raise ValueError(
            "ZIP file exceeds maximum allowed size of "
            f"{max_zip_size_bytes // (1024 * 1024)} MB."
        )

    os.makedirs(output_dir, exist_ok=True)

    logger = PipelineLogger(progress_callback)
    original_stdout = sys.stdout
    sys.stdout = logger
    try:
        start_time = datetime.now()

        print("=" * 75)
        print(" GOOGLE CONSOLIDATED INVESTIGATION REPORT GENERATOR (Web)")
        print("=" * 75)
        print(f"Source ZIP : {zip_path}")
        print(f"Output     : {output_dir}")

        working_dir = os.path.join(
            output_dir,
            "_WORKING_EXTRACTION"
        )
        os.makedirs(working_dir, exist_ok=True)

        print("\n[+] Extracting main ZIP...")
        safe_extract(zip_path, working_dir)

        print("[+] Recursively extracting nested ZIPs...")
        nested_count = extract_all_nested_zips(working_dir)
        print(f"\n[+] Nested ZIPs extracted: {nested_count}")

        print("\n[+] Reading Google data...")
        data = collect_data(working_dir)

        print("\n" + "=" * 75)
        print(" EXTRACTION SUMMARY")
        print("=" * 75)
        print(f"HTML files               : {data['html_count']}")
        print(f"CSV files                : {data['csv_count']}")
        print(f"Google IP activity       : {len(data['google_ip'])}")
        print(f"Device records           : {len(data['devices'])}")
        print(f"Phone associations       : {len(data['target_phone'])}")
        print(f"Unique formatted IPs     : {len(data['ip_details'])}")

        if data["ip_details"]:
            source_names = sorted({
                str(
                    record.get(
                        "Source",
                        record.get(
                            "Source CSV",
                            "Unknown"
                        )
                    )
                )
                for record in data["ip_details"]
            })
            print(
                "IP detail source(s)      : "
                + ", ".join(source_names)
            )
        else:
            print(
                "[WARNING] No IP detail records could be created."
            )

        print("\n[+] Converting UTC -> IST for Google IP rows...")
        data["google_ip"] = convert_google_ip_to_ist(
            data["google_ip"]
        )

        # ----------------------------------------------------
        # IP GEOLOCATION LOOKUP (optional, requires IPINFO_TOKEN)
        # ----------------------------------------------------

        ip_info = {}
        unique_ips = sorted({
            normalize_ip(rec.get("IP Address", ""))
            for rec in data["ip_details"]
            if normalize_ip(rec.get("IP Address", ""))
        })

        token_effective = ipinfo_token or os.environ.get("IPINFO_TOKEN")
        if unique_ips:
            if token_effective:
                print(
                    f"\n[+] Looking up geolocation for "
                    f"{len(unique_ips)} unique IP(s) via "
                    f"ipinfo.io..."
                )
                ip_info = lookup_ip_info(
                    unique_ips,
                    token=token_effective,
                )
                success_count = sum(
                    1 for v in ip_info.values()
                    if v and not v.get("error")
                )
                fail_count = len(ip_info) - success_count
                print(
                    f"    Geolocation results: "
                    f"{success_count} ok / {fail_count} failed"
                )
            else:
                print(
                    "\n[INFO] IPINFO_TOKEN not set; skipping "
                    "geolocation lookup. Set the IPINFO_TOKEN "
                    "environment variable to enable IP "
                    "geolocation in IP_Geolocation_Lookup.xlsx."
                )
        else:
            print(
                "\n[INFO] No unique IPs found; skipping "
                "geolocation lookup."
            )

        # ----------------------------------------------------
        # CREATE REPORTS
        # ----------------------------------------------------

        print("\n[+] Creating Excel report...")
        excel_file = create_excel_report(
            output_dir, zip_path, data, ip_info=ip_info or None
        )

        print("[+] Creating Word report...")
        docx_file = create_docx_report(
            output_dir, zip_path, data
        )

        print("[+] Creating formatted IP CSV...")
        csv_file = create_ip_csv(
            output_dir, data["ip_details"]
        )

        print("[+] Creating separate IP Word files...")
        (
            special_word_file,
            remaining_word_file,
            special_count,
            remaining_count,
        ) = create_separate_ip_word_files(
            output_dir, data["ip_details"]
        )

        print("[+] Creating IP text files...")
        (
            special_txt_file,
            remaining_txt_file,
            special_txt_count,
            remaining_txt_count,
        ) = create_ip_text_files(
            output_dir, data["ip_details"]
        )

        geo_excel_file = None
        if ip_info:
            print("[+] Creating IP geolocation Excel...")
            geo_excel_file = create_ip_geolocation_excel(
                output_dir, ip_info
            )

        elapsed = (
            datetime.now() - start_time
        ).total_seconds()

        print("\n" + "=" * 75)
        print(" REPORT GENERATION COMPLETED")
        print("=" * 75)
        print(f"Processing Time: {elapsed:.2f} seconds")
        print("Done.")
    finally:
        sys.stdout = original_stdout

    try:
        shutil.rmtree(working_dir, ignore_errors=True)
    except Exception:
        pass

    # Stats for the UI summary tiles.
    df_google_ip = pd.DataFrame(data["google_ip"])
    df_devices = pd.DataFrame(data["devices"])
    df_target = pd.DataFrame(data["target_phone"])
    df_ip_details = pd.DataFrame(data["ip_details"])

    successful = 0
    failed = 0
    unique_google_ips = 0
    if (
        not df_google_ip.empty
        and "Activity Type" in df_google_ip.columns
    ):
        activity = (
            df_google_ip["Activity Type"]
            .astype(str)
            .str.strip()
            .str.lower()
        )
        successful = int(activity.isin(
            ["login", "successful login", "success"]
        ).sum())
        failed = int(activity.isin(
            ["failed login", "failed", "failure"]
        ).sum())
        unique_google_ips = int(
            df_google_ip["IP Address"].astype(str).nunique()
        )

    connected = 0
    if (
        not df_target.empty
        and "Connection Status" in df_target.columns
    ):
        connected = int(
            df_target["Connection Status"]
            .astype(str)
            .str.strip()
            .str.upper()
            .eq("CONNECTED")
            .sum()
        )

    stats = {
        "total_ip": len(df_google_ip),
        "successful": successful,
        "failed": failed,
        "unique_google_ips": unique_google_ips,
        "formatted_ip_details": len(df_ip_details),
        "special_count": special_count,
        "remaining_count": remaining_count,
        "special_txt_count": special_txt_count,
        "remaining_txt_count": remaining_txt_count,
        "assoc_count": len(df_target),
        "connected_count": connected,
        "device_count": len(df_devices),
        "html_count": data["html_count"],
        "csv_count": data["csv_count"],
        "geolocated_count": sum(
            1 for v in ip_info.values()
            if v and not v.get("error")
        ),
        "geo_failed_count": sum(
            1 for v in ip_info.values()
            if v and v.get("error")
        ),
    }

    return {
        "excel_file": excel_file,
        "docx_file": docx_file,
        "csv_file": csv_file,
        "ipv6_word_file": special_word_file,
        "remaining_word_file": remaining_word_file,
        "ipv6_txt_file": special_txt_file,
        "remaining_txt_file": remaining_txt_file,
        "geo_excel_file": geo_excel_file,
        "stats": stats,
        "log_lines": logger.lines(),
    }


# ============================================================
# MAIN
# ============================================================

def main():

    started = datetime.now()

    try:

        # ----------------------------------------------------
        # USER INPUT
        # ----------------------------------------------------

        zip_path, output_dir = (
            get_input_paths()
        )

        print("\n" + "=" * 75)
        print(" INPUT / OUTPUT")
        print("=" * 75)

        print(
            f"Source ZIP : {zip_path}"
        )

        print(
            f"Output     : {output_dir}"
        )

        # ----------------------------------------------------
        # WORKING EXTRACTION
        # ----------------------------------------------------

        working_dir = os.path.join(
            output_dir,
            "_WORKING_EXTRACTION"
        )

        os.makedirs(
            working_dir,
            exist_ok=True
        )

        print(
            "\n[+] Extracting main ZIP..."
        )

        safe_extract(
            zip_path,
            working_dir
        )

        nested_count = (
            extract_all_nested_zips(
                working_dir
            )
        )

        print(
            f"\n[+] Nested ZIPs extracted: "
            f"{nested_count}"
        )

        # ----------------------------------------------------
        # COLLECT DATA
        # ----------------------------------------------------

        print(
            "\n[+] Reading Google data..."
        )

        data = collect_data(
            working_dir
        )

        print("\n" + "=" * 75)
        print(" EXTRACTION SUMMARY")
        print("=" * 75)

        print(
            f"HTML files               : "
            f"{data['html_count']}"
        )

        print(
            f"CSV files                : "
            f"{data['csv_count']}"
        )

        print(
            f"Google IP activity       : "
            f"{len(data['google_ip'])}"
        )

        print(
            f"Device records           : "
            f"{len(data['devices'])}"
        )

        print(
            f"Phone associations       : "
            f"{len(data['target_phone'])}"
        )

        print(
            f"Unique formatted IPs     : "
            f"{len(data['ip_details'])}"
        )

        if data["ip_details"]:
            source_names = sorted({
                str(
                    record.get(
                        "Source",
                        record.get(
                            "Source CSV",
                            "Unknown"
                        )
                    )
                )
                for record in data["ip_details"]
            })

            print(
                "IP detail source(s)      : "
                + ", ".join(source_names)
            )
        else:
            print(
                "[WARNING] No IP detail records "
                "could be created."
            )

        # ----------------------------------------------------
        # GOOGLE IP UTC -> IST
        # ----------------------------------------------------

        data["google_ip"] = (
            convert_google_ip_to_ist(
                data["google_ip"]
            )
        )

        # ----------------------------------------------------
        # CREATE REPORTS
        # ----------------------------------------------------

        print(
            "\n[+] Creating Excel report..."
        )

        excel_file = create_excel_report(
            output_dir,
            zip_path,
            data
        )

        print(
            "[+] Creating Word report..."
        )

        docx_file = create_docx_report(
            output_dir,
            zip_path,
            data
        )

        print(
            "[+] Creating formatted IP CSV..."
        )

        csv_file = create_ip_csv(
            output_dir,
            data["ip_details"]
        )

        print(
            "[+] Creating separate IP Word files..."
        )

        (
            special_word_file,
            remaining_word_file,
            special_count,
            remaining_count,
        ) = create_separate_ip_word_files(
            output_dir,
            data["ip_details"]
        )

        print(
            "[+] Creating IP text files..."
        )

        (
            special_txt_file,
            remaining_txt_file,
            special_txt_count,
            remaining_txt_count,
        ) = create_ip_text_files(
            output_dir,
            data["ip_details"]
        )

        # ----------------------------------------------------
        # REMOVE TEMP EXTRACTION
        # ----------------------------------------------------

        try:

            shutil.rmtree(
                working_dir,
                ignore_errors=True
            )

        except Exception:
            pass

        elapsed = (
            datetime.now() - started
        ).total_seconds()

        # ----------------------------------------------------
        # FINAL
        # ----------------------------------------------------

        print("\n" + "=" * 75)
        print(
            " REPORT GENERATION COMPLETED"
        )
        print("=" * 75)

        print(
            f"\nOutput Directory:\n"
            f"  {output_dir}"
        )

        print(
            f"\nExcel:\n"
            f"  {excel_file}"
        )

        print(
            f"\nWord:\n"
            f"  {docx_file}"
        )

        print(
            f"\nFormatted IP CSV:\n"
            f"  {csv_file}"
        )

        print(
            f"\n2405/2409 Word file "
            f"({special_count} IPs):\n"
            f"  {special_word_file}"
        )

        print(
            f"\nRemaining IP Word file "
            f"({remaining_count} IPs):\n"
            f"  {remaining_word_file}"
        )

        print(
            f"\n2405/2409 IP Text file "
            f"({special_txt_count} IPs):\n"
            f"  {special_txt_file}"
        )

        print(
            f"\nRemaining IP Text file "
            f"({remaining_txt_count} IPs):\n"
            f"  {remaining_txt_file}"
        )

        print(
            "\nIP Format:"
        )

        print(
            "  2405:/2409: -> IPV6 | "
            "YYYYMMDD | HHMMSS"
        )

        print(
            "  Remaining   -> IP Address | "
            "DD-MMM-YYYY HH:MM:SS"
        )

        print(
            "  FROM time = source - 1 second"
        )

        print(
            "  TO time   = source + 1 second"
        )

        print(
            "  2405:/2409: -> YYYYMMDD HHMMSS"
        )

        print(
            "  Other IPs   -> DD-MMM-YYYY HH:MM:SS"
        )

        print(
            "\nDuplicate IPs:"
        )

        print(
            "  Removed; first occurrence retained."
        )

        print(
            f"\nProcessing Time: "
            f"{elapsed:.2f} seconds"
        )

        print("\nDone.")

    except KeyboardInterrupt:

        print(
            "\n[!] Process cancelled by user."
        )

    except zipfile.BadZipFile:

        print(
            "\n[ERROR] Invalid/corrupt ZIP file."
        )

    except Exception as exc:

        print(
            "\n[ERROR]"
        )

        print(
            str(exc)
        )

        print(
            "\nPlease copy the complete error "
            "message if the problem continues."
        )


if __name__ == "__main__":
    main()